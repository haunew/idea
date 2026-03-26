#!/usr/bin/env python3
"""
Word 文档字数统计工具
支持 .docx 格式的论文字数统计
"""

import re
import sys
from pathlib import Path
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    print("错误: 请先安装 python-docx")
    print("pip install python-docx")
    sys.exit(1)


def count_words_in_text(text: str) -> dict:
    """统计文本字数"""
    # 中文字符
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    # 英文单词
    english_words = len(re.findall(r'[a-zA-Z]+', text))
    # 数字
    numbers = len(re.findall(r'\d+', text))
    # 总字符（不含空白）
    total_chars = len(re.sub(r'\s', '', text))
    
    return {
        'chinese_chars': chinese_chars,
        'english_words': english_words,
        'numbers': numbers,
        'total_chars': total_chars,
        'estimated_words': chinese_chars + english_words
    }


def analyze_docx_structure(doc: Document) -> dict:
    """分析 Word 文档结构"""
    sections = defaultdict(list)
    current_section = '前言/其他'
    
    # 章节标题关键词
    section_keywords = {
        '摘要': ['摘要', 'abstract'],
        '引言': ['引言', 'introduction', '绪论', '前言'],
        '文献综述': ['文献综述', 'literature review', '相关研究'],
        '方法': ['方法', 'methodology', 'methods', '研究方法', '实验设计'],
        '结果': ['结果', 'results', '实验结果', '研究发现'],
        '讨论': ['讨论', 'discussion'],
        '结论': ['结论', 'conclusion', '总结'],
        '参考文献': ['参考文献', 'references', 'bibliography'],
        '致谢': ['致谢', 'acknowledgements', 'acknowledgments'],
    }
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 检测是否为章节标题
        is_heading = False
        lower_text = text.lower()
        
        for section_name, keywords in section_keywords.items():
            for keyword in keywords:
                # 匹配标题：以关键词开头，且长度较短
                if (lower_text.startswith(keyword) and 
                    len(text) < 50 and 
                    (para.style.name.startswith('Heading') or 
                     any(c in text for c in ['1', '2', '一', '二', '第']))):
                    current_section = section_name
                    is_heading = True
                    break
            if is_heading:
                break
        
        if not is_heading:
            sections[current_section].append(text)
    
    return sections


def format_number(n: int) -> str:
    """格式化数字"""
    return f"{n:,}"


def main():
    if len(sys.argv) < 2:
        print("用法: python docx_word_count.py <Word文档路径>")
        print("示例: python docx_word_count.py paper.docx")
        sys.exit(1)
    
    file_path = Path(sys.argv[1])
    if not file_path.exists():
        print(f"错误: 文件不存在 {file_path}")
        sys.exit(1)
    
    if file_path.suffix.lower() != '.docx':
        print(f"错误: 仅支持 .docx 格式，当前文件: {file_path.suffix}")
        sys.exit(1)
    
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"错误: 无法读取文档 - {e}")
        sys.exit(1)
    
    # 提取所有文本
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    
    # 总体统计
    stats = count_words_in_text(full_text)
    
    print("=" * 60)
    print(f"📄 Word 文档统计报告: {file_path.name}")
    print("=" * 60)
    print(f"中文字符数: {format_number(stats['chinese_chars']):>12}")
    print(f"英文单词数: {format_number(stats['english_words']):>12}")
    print(f"数字个数:   {format_number(stats['numbers']):>12}")
    print(f"总字符数:   {format_number(stats['total_chars']):>12}")
    print("-" * 60)
    print(f"估算总词数: {format_number(stats['estimated_words']):>12}")
    print("=" * 60)
    
    # 结构分析
    sections = analyze_docx_structure(doc)
    
    if len(sections) > 1:
        print("\n📊 各章节字数分布:")
        print("-" * 60)
        print(f"{'章节':<20} {'中文字符':>12} {'英文单词':>12} {'合计':>12}")
        print("-" * 60)
        
        total_chinese = 0
        total_english = 0
        
        # 按常见顺序排序
        order = ['摘要', '引言', '文献综述', '方法', '结果', '讨论', '结论', '参考文献', '致谢', '前言/其他']
        sorted_sections = sorted(sections.keys(), 
                                key=lambda x: order.index(x) if x in order else 99)
        
        for section in sorted_sections:
            content = '\n'.join(sections[section])
            section_stats = count_words_in_text(content)
            total_chinese += section_stats['chinese_chars']
            total_english += section_stats['english_words']
            
            print(f"{section:<20} "
                  f"{format_number(section_stats['chinese_chars']):>12} "
                  f"{format_number(section_stats['english_words']):>12} "
                  f"{format_number(section_stats['estimated_words']):>12}")
        
        print("-" * 60)
        print(f"{'合计':<20} {format_number(total_chinese):>12} "
              f"{format_number(total_english):>12} "
              f"{format_number(total_chinese + total_english):>12}")
        print("=" * 60)
    
    # 表格统计
    table_count = len(doc.tables)
    if table_count > 0:
        print(f"\n📋 表格数量: {table_count}")
    
    # 图片统计（通过段落中的 inline shapes 估算）
    image_count = 0
    for para in doc.paragraphs:
        image_count += len(para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
    
    if image_count > 0:
        print(f"🖼️  图片数量: {image_count}")


if __name__ == '__main__':
    main()
